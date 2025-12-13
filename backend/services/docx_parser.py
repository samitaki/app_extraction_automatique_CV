from docx import Document

## Dans le cas des documents docx , on prend en considereation le contenu des paragraphes et des tableaux invisibles
def parse_docx(file_path: str) -> str:
    doc = Document(file_path)
    # Récupérer tous les paragraphes
    paragraphs = [p.text for p in doc.paragraphs]

    table_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text.append(cell.text)

    # Combiner tout le texte
    full_text = "\n".join(paragraphs + table_text)

    return full_text